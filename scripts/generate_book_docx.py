#!/usr/bin/env python3
"""Generate a book-style DOCX booklet from the AI book launch draft."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SOURCE = ROOT / "docs" / "ai_book_business" / "08_validation_handbook_full_draft.md"
DEFAULT_OUTPUT = ROOT / "output" / "doc" / "7天做出100条短视频脚本_数字验证款首版.docx"


def set_run_font(run, *, name: str, size: int, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_style(
    paragraph,
    *,
    before: int = 0,
    after: int = 0,
    line_spacing: float = 1.6,
    first_line_indent_cm: float = 0.74,
    left_indent_cm: float = 0,
    right_indent_cm: float = 0,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.first_line_indent = Cm(first_line_indent_cm) if first_line_indent_cm else None
    fmt.left_indent = Cm(left_indent_cm) if left_indent_cm else None
    fmt.right_indent = Cm(right_indent_cm) if right_indent_cm else None
    paragraph.alignment = alignment


def shade_paragraph(paragraph, fill: str = "F3F4F6") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, name="SimSun", size=9)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_paragraph.add_run(title)
    set_run_font(header_run, name="SimSun", size=9)

    footer_paragraph = section.footer.paragraphs[0]
    add_page_number(footer_paragraph)


def extract_structure(lines: list[str]) -> tuple[str, str, list[str], list[str], bool]:
    title = "7 天做出 100 条短视频脚本"
    subtitle = "用 DeepSeek 批量写口播、选题、评论区钩子与成交结尾"
    body_lines: list[str] = []
    toc_items: list[str] = []
    has_preface = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            continue
        if stripped.startswith("副标题:"):
            subtitle = stripped.split(":", 1)[1].strip()
            continue
        if stripped.startswith("## "):
            toc_items.append(stripped[3:].strip())
        elif stripped:
            has_preface = True
        body_lines.append(line)

    return title, subtitle, body_lines, toc_items, has_preface


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(130)
    title_paragraph.paragraph_format.space_after = Pt(16)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, name="Microsoft YaHei", size=24, bold=True)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    set_run_font(subtitle_run, name="Microsoft YaHei", size=12)

    meta_paragraph = doc.add_paragraph()
    meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_paragraph.paragraph_format.space_after = Pt(8)
    meta_run = meta_paragraph.add_run("19.9 元数字验证款首版手册")
    set_run_font(meta_run, name="SimSun", size=11)

    note_paragraph = doc.add_paragraph()
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_paragraph.add_run("用于首轮验证交付，不作为正式出版扉页信息")
    set_run_font(note_run, name="SimSun", size=9)

    doc.add_page_break()


def add_contents_page(doc: Document, toc_items: list[str], has_preface: bool) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("目录")
    set_run_font(run, name="Microsoft YaHei", size=18, bold=True)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(18)

    items = ["写在前面", *toc_items] if has_preface else toc_items

    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        set_paragraph_style(
            paragraph,
            after=4,
            line_spacing=1.35,
            first_line_indent_cm=0,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        label = f"{index:02d}. {item}"
        run = paragraph.add_run(label)
        set_run_font(run, name="SimSun", size=11)

    note = doc.add_paragraph()
    set_paragraph_style(
        note,
        before=10,
        after=0,
        line_spacing=1.25,
        first_line_indent_cm=0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    note_run = note.add_run("建议在 Word 中全选文档后再微调页眉、页脚和分页。")
    set_run_font(note_run, name="SimSun", size=9)

    doc.add_page_break()


def add_major_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(16)
    run = paragraph.add_run(text)
    set_run_font(run, name="Microsoft YaHei", size=18, bold=True)


def add_minor_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_style(
        paragraph,
        before=10,
        after=6,
        line_spacing=1.2,
        first_line_indent_cm=0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run(text)
    set_run_font(run, name="Microsoft YaHei", size=13, bold=True)


def add_lead_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_style(
        paragraph,
        before=6,
        after=4,
        line_spacing=1.3,
        first_line_indent_cm=0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run(text.rstrip(":"))
    set_run_font(run, name="Microsoft YaHei", size=11, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_style(paragraph, after=4, line_spacing=1.68)
    run = paragraph.add_run(text)
    set_run_font(run, name="SimSun", size=11)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_style(
        paragraph,
        after=3,
        line_spacing=1.35,
        first_line_indent_cm=0,
        left_indent_cm=0.4,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run(text)
    set_run_font(run, name="SimSun", size=10.5)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_style(
        paragraph,
        after=3,
        line_spacing=1.35,
        first_line_indent_cm=0,
        left_indent_cm=0.4,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run(text)
    set_run_font(run, name="SimSun", size=10.5)


def add_code_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_style(
        paragraph,
        before=1,
        after=1,
        line_spacing=1.15,
        first_line_indent_cm=0,
        left_indent_cm=0.7,
        right_indent_cm=0.4,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run(text)
    set_run_font(run, name="Menlo", size=9)
    shade_paragraph(paragraph)


def render_body(doc: Document, lines: list[str]) -> None:
    in_code_block = False
    seen_major_heading = False
    preface_heading_added = False

    for raw_line in lines:
        stripped = raw_line.strip()

        if stripped == "```":
            in_code_block = not in_code_block
            continue

        if not stripped:
            continue

        if in_code_block:
            add_code_line(doc, raw_line.rstrip("\n"))
            continue

        if stripped.startswith("## "):
            if seen_major_heading:
                doc.add_page_break()
            add_major_heading(doc, stripped[3:].strip())
            seen_major_heading = True
            continue

        if stripped.startswith("### "):
            if not seen_major_heading and not preface_heading_added:
                add_major_heading(doc, "写在前面")
                preface_heading_added = True
            add_minor_heading(doc, stripped[4:].strip())
            continue

        if stripped.startswith("- "):
            if not seen_major_heading and not preface_heading_added:
                add_major_heading(doc, "写在前面")
                preface_heading_added = True
            add_bullet(doc, stripped[2:].strip())
            continue

        numbered = stripped.split(". ", 1)
        if len(numbered) == 2 and numbered[0].isdigit():
            if not seen_major_heading and not preface_heading_added:
                add_major_heading(doc, "写在前面")
                preface_heading_added = True
            add_numbered(doc, numbered[1].strip())
            continue

        if stripped.endswith(":") and len(stripped) <= 18:
            if not seen_major_heading and not preface_heading_added:
                add_major_heading(doc, "写在前面")
                preface_heading_added = True
            add_lead_line(doc, stripped)
            continue

        if not seen_major_heading and not preface_heading_added:
            add_major_heading(doc, "写在前面")
            preface_heading_added = True
        add_body_paragraph(doc, stripped)


def build_docx(source_path: Path, output_path: Path) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    title, subtitle, body_lines, toc_items, has_preface = extract_structure(lines)

    doc = Document()
    configure_document(doc, title)
    add_cover(doc, title, subtitle)
    add_contents_page(doc, toc_items, has_preface)
    render_body(doc, body_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate book-style DOCX from markdown draft.")
    parser.add_argument("source", nargs="?", default=str(DEFAULT_SOURCE))
    parser.add_argument("output", nargs="?", default=str(DEFAULT_OUTPUT))
    args = parser.parse_args()

    build_docx(Path(args.source), Path(args.output))
    print(f"Wrote {Path(args.output)}")


if __name__ == "__main__":
    main()
